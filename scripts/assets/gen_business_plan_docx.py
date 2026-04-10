# -*- coding: utf-8 -*-
"""Render a Markdown business plan to DOCX."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"
DEFAULT_SRC = DOCS / "商业计划书-BrightAgent-融资比赛版.md"
DEFAULT_DST = DOCS / "business-plan-brightagent-full.docx"

NAVY = RGBColor(30, 58, 95)
ACCENT = RGBColor(37, 99, 235)
MUTED = RGBColor(99, 115, 129)


def strip_md(text: str) -> str:
    text = re.sub(r"^>\s*", "", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def parse_table_row(line: str) -> list[str]:
    return [strip_md(cell.strip()) for cell in line.strip().strip("|").split("|")]


def is_separator_row(line: str) -> bool:
    return bool(re.fullmatch(r"\|\s*[-:| ]+\|", line.strip()))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_run(run, *, size: int | None = None, bold: bool = False, color=None, font_name: str = "Microsoft YaHei") -> None:
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("BrightAgent")
    format_run(run, size=24, bold=True, color=NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    run = p2.add_run("Agent Marketplace & Trust Layer")
    format_run(run, size=11, color=MUTED, font_name="Calibri")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(26)
    p3.paragraph_format.space_after = Pt(8)
    run = p3.add_run(title)
    format_run(run, size=20, bold=True, color=NAVY)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(18)
    run = p4.add_run(subtitle)
    format_run(run, size=11, color=MUTED)

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run("_" * 46)
    format_run(run, size=12, color=ACCENT, font_name="Calibri")

    doc.add_page_break()


def main() -> None:
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SRC
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_DST

    if not src.is_absolute():
        src = ROOT / src
    if not dst.is_absolute():
        dst = ROOT / dst

    if not src.is_file():
        raise SystemExit(f"Missing {src}")

    raw_lines = src.read_text(encoding="utf-8").splitlines()
    lines = [line.rstrip() for line in raw_lines]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(11)

    title = "BrightAgent 商业计划书"
    subtitle = "融资 / 比赛展示版"
    for i, line in enumerate(lines[:6]):
        if line.startswith("# "):
            title = strip_md(line[2:])
        elif line.startswith("> "):
            text = strip_md(line)
            if "Agent as Service" in text:
                subtitle = text

    add_cover(doc, title, subtitle)

    i = 0
    while i < len(lines):
        raw = lines[i]
        cleaned = strip_md(raw)

        if raw.startswith("# "):
            i += 1
            continue
        if raw.strip() == "---":
            i += 1
            continue
        if not cleaned:
            i += 1
            continue
        if raw.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(cleaned)
            format_run(run, size=10, color=MUTED)
            i += 1
            continue
        if raw.lstrip().startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = [parse_table_row(line) for line in table_lines if not is_separator_row(line)]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                table.autofit = True
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        format_run(run, size=10, bold=(r_idx == 0), color=(NAVY if r_idx == 0 else BODY_COLOR))
                        if r_idx == 0:
                            set_cell_shading(cell, "EBF1FF")
                doc.add_paragraph()
            continue
        if cleaned.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(cleaned[3:].strip())
            format_run(run, size=15, bold=True, color=NAVY)
            i += 1
            continue
        if cleaned.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(cleaned[4:].strip())
            format_run(run, size=12, bold=True, color=ACCENT)
            i += 1
            continue
        if re.match(r"^\d+\.\s", cleaned):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(re.sub(r"^\d+\.\s*", "", cleaned))
            format_run(run, size=11)
            i += 1
            continue
        if re.match(r"^[-*]\s", cleaned):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(re.sub(r"^[-*]\s*", "", cleaned))
            format_run(run, size=11)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(cleaned)
        format_run(run, size=11)
        i += 1

    # Add page numbers to footer.
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("BrightAgent")
        format_run(run, size=9, color=MUTED)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    print(f"Wrote {dst}")


BODY_COLOR = RGBColor(45, 55, 65)


if __name__ == "__main__":
    main()
