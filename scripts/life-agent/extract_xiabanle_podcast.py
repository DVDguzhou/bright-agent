# -*- coding: utf-8 -*-
"""Extract guest/host transcript from 《我下班了》 docx files."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

DOWNLOADS = Path(r"c:\Users\Caiqing\Downloads")
OUT = Path(__file__).resolve().parent / "xiabanle_episodes.json"

# guest sp = 逐字稿中的发言人编号；另一编号为主持人阿拉
EPISODES = [
    {"ep": "03", "find": ["03", "成都土著"], "guest": "文婷", "aliases": [], "sp": 1,
     "title": "成都土著小姐姐，主业B端运营，副业开中年女性家居服店", "display": "成都土著家居服"},
    {"ep": "04", "find": ["04", "躺平"], "guest": "阿基", "aliases": ["明宇", "阿宇"], "sp": 1,
     "title": "这批来成都躺平的90后，都卷起来创业了", "display": "成都躺平创业"},
    {"ep": "05", "find": ["05", "闯江湖"], "guest": "范寒", "aliases": ["蒋成池"], "sp": 1,
     "title": "在一线城市闯江湖，95后回二线打副本", "display": "一线闯二线副本"},
    {"ep": "06", "find": ["06", "社交app"], "guest": "黄豆老师、小鱼儿", "aliases": [], "sp": "dual",
     "title": "男女生刷社交app，发展一种很新的亲密关系", "display": "社交app新关系"},
    {"ep": "07", "find": ["07", "7份工作"], "guest": "小郭", "aliases": ["郭荣"], "sp": 1,
     "title": "8年换7份工作，找工作不要病急乱投医", "display": "七年七份工作"},
    {"ep": "08", "find": ["008", "盖洛普"], "guest": "大鱼", "aliases": [], "sp": "unsplit",
     "title": "产品经理转化盖洛普优势咨询，自由职业太美好了", "display": "盖洛普自由职业"},
    {"ep": "09", "find": ["009", "热爱生活"], "guest": "狗哥", "aliases": ["杨敏敏"], "sp": 1,
     "title": "人到30，从热爱工作转向热爱生活", "display": "30岁热爱生活"},
    {"ep": "10", "find": ["010", "字节裁员"], "guest": "陈楠", "aliases": [], "sp": 1,
     "title": "31岁，刚年薪百万就被字节裁员，什么感受？", "display": "字节裁员百万"},
    {"ep": "11", "find": ["011", "Gap2"], "guest": "耗子学长", "aliases": [], "sp": 2,
     "title": "毕业6年Gap2年，探索AI+自媒体创业之路", "display": "Gap探索AI创业"},
    {"ep": "12", "find": ["012", "Sina"], "guest": "思娜", "aliases": ["Sina"], "sp": 2,
     "title": "独立运营顾问Sina，把热爱变成事业", "display": "运营顾问Sina"},
    {"ep": "13", "find": ["013", "B面人生"], "guest": "七天", "aliases": [], "sp": 2,
     "title": "在杭州，帮1000+年轻人探索B面人生", "display": "杭州B面人生"},
    {"ep": "14", "find": ["014", "成都找工作"], "guest": "范范、Peter、君明", "aliases": [], "sp": "multi",
     "title": "成都找工作有点具体！薪资曝光，人均自由职业不上班！", "display": "成都自由职业"},
    {"ep": "15", "find": ["015", "卖水果"], "guest": "红伶", "aliases": ["红林"], "sp": 2,
     "title": "前阿里运营回成都卖水果！要创业更要快乐！", "display": "阿里回成都卖水果"},
    {"ep": "16", "find": ["016", "旅游博主"], "guest": "苏苏", "aliases": [], "sp": 2,
     "title": "什么旅游博主能月入大几万？商业模式坦白局", "display": "旅游博主月入几万"},
    {"ep": "17", "find": ["017", "辍学"], "guest": "于小航", "aliases": ["于航"], "sp": 2,
     "title": "辍学4年重回校园，把爱好做成百万影视公司！", "display": "辍学百万影视"},
]

HOST_MARKERS = ("我是阿拉", "主播阿拉", "欢迎来到播客，我下班了", "欢迎来到播客 我下班了", "我下班了，我是")


def find_file(keys: list[str]) -> Path | None:
    for fn in DOWNLOADS.iterdir():
        if fn.suffix != ".docx":
            continue
        name = fn.name
        if all(k in name for k in keys[:1]) and any(k in name for k in keys):
            return fn
    return None


def parse_numbered(path: Path) -> tuple[str, dict[int, list[str]]]:
    doc = Document(str(path))
    speaker_re = re.compile(r"^发言人(\d+)\s+\d{2}:\d{2}\s*$")
    current = 0
    buckets: dict[int, list[str]] = {1: [], 2: []}
    meta = ""
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        m = speaker_re.match(t)
        if m:
            current = int(m.group(1))
            continue
        if current in buckets:
            buckets[current].append(t)
        elif not meta and not re.match(r"^\d{4}", t):
            meta = t
    return meta, buckets


def parse_unsplit(path: Path) -> tuple[str, list[str], list[str]]:
    """EP08: 发言人 without number — alternate guest/host blocks."""
    doc = Document(str(path))
    marker_re = re.compile(r"^发言人\s*\d{2}:\d{2}\s*$")
    blocks: list[list[str]] = []
    cur: list[str] = []
    meta = ""
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or re.match(r"^\d{4}", t):
            continue
        if marker_re.match(t) or (t.startswith("发言人") and re.search(r"\d{2}:\d{2}", t)):
            if cur:
                blocks.append(cur)
            cur = []
            continue
        if not meta and not t.startswith("发言人"):
            meta = t
            continue
        cur.append(t)
    if cur:
        blocks.append(cur)
    guest, host = [], []
    for i, blk in enumerate(blocks):
        text = "\n".join(blk)
        if i % 2 == 0:
            guest.extend(blk)
        else:
            host.extend(blk)
    if not guest and blocks:
        guest = blocks[0]
        host = sum(blocks[1:], [])
    return meta, guest, host


def is_host_line(t: str) -> bool:
    head = t[:80]
    return any(m in head for m in HOST_MARKERS) or (head.startswith("Hello") and "阿拉" in t[:120])


def build_dual(buckets: dict[int, list[str]]) -> tuple[str, str]:
    g1 = "\n".join(buckets.get(1, []))
    g2 = "\n".join(buckets.get(2, []))
    guest = f"【黄豆老师】\n{g1}\n\n【小鱼儿】\n{g2}".strip()
    host_lines = [t for t in buckets.get(2, []) if is_host_line(t)]
    host = "\n".join(host_lines[:40])
    return guest, host


def build_multi(buckets: dict[int, list[str]]) -> tuple[str, str]:
    g1 = "\n".join(buckets.get(1, []))
    g2 = "\n".join(buckets.get(2, []))
    guest = f"【范范等嘉宾】\n{g1}\n\n【Peter / 君明等】\n{g2}".strip()
    host = "\n".join(t for t in buckets.get(1, []) + buckets.get(2, []) if is_host_line(t))[:6000]
    return guest, host


def build_knowledge(ep: dict, meta: str, guest: str, host: str) -> str:
    alias = f"（{' / '.join(ep['aliases'])}）" if ep["aliases"] else ""
    parts = [
        f"节目：《我下班了》播客 EP{ep['ep']}《{ep['title']}》",
        f"嘉宾：{ep['guest']}{alias}",
        "主持人：阿拉",
        "",
        "【嘉宾分享】",
        guest.strip(),
    ]
    if host.strip():
        parts += ["", "【主持人追问与补充（供理解语境）】", host.strip()[:8000]]
    return "\n".join(parts).strip()


def first_sentence(text: str, max_len: int = 140) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return ""
    for sep in "。！？\n":
        idx = text.find(sep)
        if 20 < idx < max_len:
            return text[: idx + 1]
    return text[:max_len] + ("…" if len(text) > max_len else "")


def main() -> None:
    profiles = []
    for ep in EPISODES:
        path = find_file(ep["find"])
        if not path:
            print("MISSING", ep["ep"], ep["find"])
            continue
        sp = ep["sp"]
        if sp == "unsplit":
            meta, guest_lines, host_lines = parse_unsplit(path)
            guest_text = "\n".join(guest_lines)
            host_text = "\n".join(host_lines)
        else:
            meta, buckets = parse_numbered(path)
            if sp == "dual":
                guest_text, host_text = build_dual(buckets)
            elif sp == "multi":
                guest_text, host_text = build_multi(buckets)
            else:
                guest_sp = int(sp)
                host_sp = 2 if guest_sp == 1 else 1
                guest_text = "\n".join(buckets.get(guest_sp, []))
                host_text = "\n".join(buckets.get(host_sp, []))

        kb = build_knowledge(ep, meta, guest_text, host_text)
        short = first_sentence(guest_text) or ep["title"]
        profiles.append({
            **ep,
            "path": str(path),
            "meta": meta,
            "short_bio": short,
            "knowledge_body": kb,
            "guest_len": len(guest_text),
            "host_len": len(host_text),
        })
        print(ep["ep"], ep["display"], profiles[-1]["guest_len"], profiles[-1]["host_len"])

    OUT.write_text(json.dumps(profiles, ensure_ascii=False, indent=2), encoding="utf-8")
    print("wrote", len(profiles), "profiles ->", OUT)


if __name__ == "__main__":
    main()
